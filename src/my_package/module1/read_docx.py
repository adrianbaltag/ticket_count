# """This module reads a Word document and returns the first, 2nd, and last paragraphs"""

# import os

# from colorama import Fore
# from docx import Document


# def get_last_non_empty_paragraph(doc):
#     """Get the last non-empty paragraph in a Word document"""
#     return next(
#         (para.text for para in reversed(doc.paragraphs) if para.text.strip()), ""
#     )


# async def read_docx(file_name):
#     """Read a Word document and return the first, 2nd, and last paragraphs"""
#     file_path = os.path.join(os.path.expanduser("~"), "Desktop", file_name)

#     try:
#         doc = Document(file_path)  # Open the document

#         # Extract first non-empty paragraph
#         nrb_ticket = next((p.text for p in doc.paragraphs if p.text.strip()), "")

#         # Extract second non-empty paragraph (skip the first one)
#         phone_number = next(
#             (p.text for i, p in enumerate(doc.paragraphs) if p.text.strip() and i > 0),
#             "",
#         )

#         # Extract last non-empty paragraph
#         ticket_status = get_last_non_empty_paragraph(doc)

#         # Print only relevant data
#         print(f"First paragraph (NRB Ticket): {nrb_ticket}")
#         print(f"Second paragraph (Phone Number): {phone_number}")
#         print(f"Last paragraph (Ticket Status): {ticket_status}")

#         return nrb_ticket, phone_number, ticket_status.lower()

#     except Exception as e:
#         print(Fore.RED + f"Error reading {file_name}: {e}")
#         return "", "", ""


"""This module reads a Word document and returns the first, 2nd, and last paragraphs"""

import os
import re

from colorama import Fore
from docx import Document


def get_last_non_empty_paragraph(doc):
    """Get the last non-empty paragraph in a Word document"""
    return next(
        (para.text for para in reversed(doc.paragraphs) if para.text.strip()), ""
    )


async def read_docx(file_name):
    """Read a Word document and return the first, 2nd, and last paragraphs"""
    file_path = os.path.join(os.path.expanduser("~"), "Desktop", file_name)

    try:
        doc = Document(file_path)  # Open the document

        # Extract first non-empty paragraph
        nrb_ticket = next((p.text for p in doc.paragraphs if p.text.strip()), "")

        # Extract second non-empty paragraph (skip the first one)
        phone_number = next(
            (p.text for i, p in enumerate(doc.paragraphs) if p.text.strip() and i > 0),
            "",
        )

        # Extract last non-empty paragraph
        last_paragraph = get_last_non_empty_paragraph(doc)

        # Extract only "wip" or "resolved" from the last paragraph
        # This is the critical fix - we only want the status word
        status_match = re.search(r"\b(wip|resolved)\b", last_paragraph.lower())
        ticket_status = (
            status_match.group(1) if status_match else "wip"
        )  # Default to "wip" if not found

        # Print only relevant data
        print(f"First paragraph (NRB Ticket): {nrb_ticket}")
        print(f"Second paragraph (Phone Number): {phone_number}")
        print(f"Last paragraph (Ticket Status): {ticket_status}")

        return nrb_ticket, phone_number, ticket_status

    except Exception as e:
        print(Fore.RED + f"Error reading {file_name}: {e}")
        return "", "", ""
